#!/usr/bin/env python3
"""
=============================================================================
SCRIPT NAME: generate_drafting_magic_synthetic_docs.py
=============================================================================

DESCRIPTION:
    Generates a synthetic estate-planning document packet for testing the
    Drafting Magic browser-side document extractor. All documents use fake
    but realistic names, addresses, and account-like facts so the privacy
    filter, extraction pipeline, comparison matrix, and review workflow can
    be exercised without real client data.

    The packet includes a revocable living trust, pour-over will, advance
    health care directive, durable financial power of attorney, prenuptial
    agreement, attorney instruction text, signing memo, a legacy .doc format
    test file, and an image-only scanned PDF test file. The documents
    intentionally contain cross-document conflicts (trust name/date mismatch,
    agent-order mismatches, property-character ambiguity, incapacity trigger
    mismatches, old HIPAA wording) to verify that Drafting Magic flags them
    for attorney review.

    Output files are written to a clean directory then zipped together.

INPUT FILES:
    /System/Library/Fonts/Supplemental/Times New Roman Bold.ttf
        System font used when rendering the image-only PDF. Falls back to
        Pillow default if unavailable (non-fatal).
    /System/Library/Fonts/Supplemental/Times New Roman.ttf
        System font used when rendering the image-only PDF. Falls back to
        Pillow default if unavailable (non-fatal).

OUTPUT FILES:
    .../fixtures/drafting-magic-synthetic-packet/00_README_TESTING_GUIDE.md
        Markdown testing guide describing each file, expected behaviors,
        suggested upload mapping, and known cross-document issues.
    .../fixtures/drafting-magic-synthetic-packet/01_REVOCABLE_LIVING_TRUST__Chen_Family_2016.docx
        Synthetic Chen Family Revocable Trust with trustee order, funding
        schedule, property-character ambiguity, and incapacity trigger.
    .../fixtures/drafting-magic-synthetic-packet/02_POUR_OVER_WILL__Maya_Chen_Trust_Date_Mismatch.pdf
        Selectable PDF pour-over will with intentional trust name/date
        mismatch against the trust document.
    .../fixtures/drafting-magic-synthetic-packet/03_ADVANCE_HEALTH_CARE_DIRECTIVE__Priya_First_Agent.docx
        Advance health care directive naming Priya Shah as first agent,
        using old HIPAA/privacy wording, with a health-care-decision
        incapacity trigger.
    .../fixtures/drafting-magic-synthetic-packet/04_DURABLE_FINANCIAL_POA__Daniel_First_Agent.docx
        Durable financial POA effective immediately, naming Daniel Chen
        first, with gifting and transfer limits.
    .../fixtures/drafting-magic-synthetic-packet/05_PRENUPTIAL_AGREEMENT__Separate_Property_Constraints.pdf
        Selectable PDF prenuptial agreement preserving separate/community
        property classifications, spousal waivers, and transfer limits.
    .../fixtures/drafting-magic-synthetic-packet/06_NEW_LAW_AND_ATTORNEY_INSTRUCTION__Packet_Reconciliation.txt
        Plain text simulating a new California packet-consistency law and
        attorney instruction for the review workflow.
    .../fixtures/drafting-magic-synthetic-packet/07_SIGNING_MEMO__Conflicting_Execution_Checklist.md
        Markdown signing memo whose claims intentionally contradict the
        packet, designed to trigger stale-matrix regeneration.
    .../fixtures/drafting-magic-synthetic-packet/08_LEGACY_DOC_FORMAT__Unsupported_Extraction.doc
        Plain text with a .doc extension to verify the extractor produces
        a "Needs review" warning for unsupported legacy Word binary files.
    .../fixtures/drafting-magic-synthetic-packet/09_SCANNED_PDF_WARNING__Image_Only_Addendum.pdf
        Image-only PDF (no selectable text layer) to verify the extractor
        marks scanned PDFs as needing review.
    .../fixtures/drafting-magic-synthetic-packet/manifest.json
        JSON manifest describing the generation tool, synthetic flag, file
        listing, and expected feature-coverage categories.
    .../fixtures/drafting-magic-synthetic-packet.zip
        Zip archive containing all of the above files for convenient
        download or upload as a single artifact.
    stdout
        Prints the absolute paths of the output directory and zip archive.

VERSION: 1.0
LAST UPDATED: 2026-06-05
AUTHOR: Arjun Divecha

DEPENDENCIES:
    python-docx
    Pillow
    reportlab

USAGE:
    python generate_drafting_magic_synthetic_docs.py

NOTES:
    - All documents are synthetic and must never be used with real client data.
    - The output directory is cleaned (all files deleted) before each run.
    - An intermediate PNG (_scanned_addendum_source.png) is created then
      deleted during the image-only PDF generation.
    - Bloomberg / market data is not used; this script runs fully offline.
=============================================================================
"""

from __future__ import annotations

import json
import textwrap
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as ReportLabImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "fixtures" / "drafting-magic-synthetic-packet"


def wrap(text: str) -> str:
    return "\n".join(line.strip() for line in textwrap.dedent(text).strip().splitlines())


def blocks_to_plain(blocks: list[tuple[str, str]]) -> str:
    lines: list[str] = []
    for title, body in blocks:
        lines.append(title.upper())
        lines.append(body)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def make_docx(path: Path, title: str, subtitle: str, blocks: list[tuple[str, str]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(8)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SYNTHETIC TEST DOCUMENT - NOT A REAL CLIENT FILE")
    r.bold = True
    r.font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(10)

    doc.add_paragraph("")
    for heading, body in blocks:
        doc.add_heading(heading, level=2)
        for para in body.split("\n\n"):
            doc.add_paragraph(para.strip())

    doc.core_properties.author = "Drafting Magic Synthetic Fixture"
    doc.core_properties.title = title
    doc.save(path)


def make_pdf(path: Path, title: str, subtitle: str, blocks: list[tuple[str, str]]) -> None:
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=title,
        author="Drafting Magic Synthetic Fixture",
    )
    base = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "FixtureTitle",
        parent=base["Title"],
        fontName="Times-Bold",
        fontSize=15,
        leading=18,
        alignment=1,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "FixtureSubtitle",
        parent=base["BodyText"],
        fontName="Times-Italic",
        fontSize=9,
        leading=12,
        alignment=1,
        spaceAfter=18,
    )
    heading_style = ParagraphStyle(
        "FixtureHeading",
        parent=base["Heading2"],
        fontName="Times-Bold",
        fontSize=11,
        leading=14,
        spaceBefore=10,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "FixtureBody",
        parent=base["BodyText"],
        fontName="Times-Roman",
        fontSize=10.5,
        leading=14,
        spaceAfter=8,
    )

    story = [
        Paragraph("SYNTHETIC TEST DOCUMENT - NOT A REAL CLIENT FILE", subtitle_style),
        Paragraph(title, title_style),
        Paragraph(subtitle, subtitle_style),
    ]
    for heading, body in blocks:
        story.append(Paragraph(heading, heading_style))
        for para in body.split("\n\n"):
            story.append(Paragraph(para.strip(), body_style))
        story.append(Spacer(1, 3))
    doc.build(story)


def make_image_only_pdf(path: Path) -> None:
    image_path = OUT / "_scanned_addendum_source.png"
    img = Image.new("RGB", (1600, 2100), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf", 54)
        font_body = ImageFont.truetype("/System/Library/Fonts/Supplemental/Times New Roman.ttf", 34)
    except Exception:
        font_title = ImageFont.load_default()
        font_body = ImageFont.load_default()

    y = 140
    draw.text((120, y), "SYNTHETIC SCANNED ADDENDUM", fill="black", font=font_title)
    y += 90
    lines = [
        "Image-only PDF for OCR warning test.",
        "",
        "This page intentionally has no selectable text layer.",
        "Drafting Magic should mark the PDF as needing review",
        "because the browser extractor cannot read scanned PDFs yet.",
        "",
        "Visible content mentions separate property, gifting,",
        "trust funding, and health care authority, but those terms",
        "should not be extracted until OCR exists.",
    ]
    for line in lines:
        draw.text((120, y), line, fill="black", font=font_body)
        y += 54
    img.save(image_path)

    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=0.5 * inch,
        leftMargin=0.5 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )
    story = [ReportLabImage(str(image_path), width=7.0 * inch, height=9.2 * inch)]
    doc.build(story)
    image_path.unlink(missing_ok=True)


TRUST_BLOCKS = [
    (
        "Article One. Trust identity and family context",
        wrap(
            """
            The Chen Family Revocable Trust is dated May 12, 2016, as amended on
            September 3, 2022. The settlors are Maya Chen and Daniel Chen, spouses,
            both residing at 4189 Cedar Harbor Lane, Palo Alto, California 94303.
            This synthetic trust uses real-looking names and addresses only to test
            the local privacy filter.
            """
        ),
    ),
    (
        "Article Two. Successor trustee appointments",
        wrap(
            """
            Maya Chen shall act as initial trustee while living and competent.
            Upon Maya Chen's incapacity or death, Daniel Chen shall serve as first
            successor trustee. If Daniel Chen cannot serve, Priya Shah shall serve
            as backup successor trustee. The trustee is a fiduciary and must act
            solely for the beneficiaries.
            """
        ),
    ),
    (
        "Article Three. Distributions and health support",
        wrap(
            """
            During incapacity, the trustee may pay health, support, maintenance,
            education, placement, in-home care, and medical expenses for the
            surviving settlor and dependent beneficiaries. The trustee may rely on
            a HIPAA release or health care agent certification when coordinating
            care expenses.
            """
        ),
    ),
    (
        "Article Four. Funding schedule and property character",
        wrap(
            """
            Schedule A funds community property and separate property into one
            administrative schedule without a property-character legend. The 2020
            Solara AI Inc. founder shares, the Palo Alto residence, and the
            Westlake brokerage account are all listed together. No clause states
            whether trust funding changes separate property, community property,
            or reimbursement rights.
            """
        ),
    ),
    (
        "Article Five. Incapacity trigger",
        wrap(
            """
            Incapacity requires written certification by two licensed physicians
            or a court determination. The trustee may also treat a settlor as
            incapacitated if the settlor cannot manage financial affairs for
            thirty consecutive days.
            """
        ),
    ),
]

WILL_BLOCKS = [
    (
        "Article One. Pour-over residue",
        wrap(
            """
            I, Maya Chen, leave the residue of my estate to the trustee of the
            Maya Chen Living Trust dated May 21, 2016, including later amendments
            and restatements. This trust name and trust date intentionally do not
            match the trust instrument in the packet.
            """
        ),
    ),
    (
        "Article Two. Executor appointments",
        wrap(
            """
            Daniel Chen is nominated as executor. If Daniel Chen cannot serve,
            Maya Chen's sister Elena Park is nominated as alternate executor. The
            executor may transfer property to the trustee after death and may
            execute assignments needed to complete the pour-over.
            """
        ),
    ),
    (
        "Article Three. Property and prenup silence",
        wrap(
            """
            The pour-over will does not restate separate property funding
            limitations, community property classification, reimbursement rights,
            or prenuptial agreement transfer constraints. It says all residue
            should pass to the trust without a property-character legend.
            """
        ),
    ),
]

AHCD_BLOCKS = [
    (
        "Section One. Health care agent order",
        wrap(
            """
            Maya Chen names herself as principal and appoints Maya Chen's spouse,
            Daniel Chen, as alternate only if the first agent cannot serve. The
            first health care agent is Priya Shah. This creates an intentional
            agent-order mismatch with the financial power of attorney.
            """
        ),
    ),
    (
        "Section Two. Treatment and placement authority",
        wrap(
            """
            The health care agent may make treatment, placement, hospice, pain
            management, end-of-life, anatomical gift, and medical transport
            decisions when Maya Chen cannot make health care decisions. The agent
            may consult with physicians and care managers.
            """
        ),
    ),
    (
        "Section Three. Privacy release",
        wrap(
            """
            The privacy release authorizes disclosure of medical information, but
            uses older HIPAA wording and does not mention digital health portals,
            telehealth accounts, or cloud-based pharmacy records.
            """
        ),
    ),
]

POA_BLOCKS = [
    (
        "Section One. Financial agent order",
        wrap(
            """
            Daniel Chen is appointed as first financial agent. Elena Park is
            appointed as alternate agent. Priya Shah is not named as financial
            agent. This differs from the advance health care directive and should
            be flagged for attorney review.
            """
        ),
    ),
    (
        "Section Two. Effective date and incapacity",
        wrap(
            """
            This durable financial power of attorney is effective immediately and
            remains effective during incapacity. No physician certification is
            required before the agent may manage banking, insurance, benefits,
            tax, retirement, real property, or digital asset matters.
            """
        ),
    ),
    (
        "Section Three. Gifting and transfers",
        wrap(
            """
            The agent may make annual exclusion gifts to descendants, but may not
            transfer separate property, community property, or trust funding assets
            in a way that expands spousal waivers without express written consent.
            The agent must honor fiduciary duties.
            """
        ),
    ),
]

PRENUP_BLOCKS = [
    (
        "Article One. Separate property classifications",
        wrap(
            """
            Maya Chen's premarital Solara AI Inc. founder shares, appreciation,
            dividends, and sale proceeds remain Maya Chen's separate property.
            Daniel Chen's Westlake brokerage account remains Daniel Chen's
            separate property. The Palo Alto residence is community property
            unless a signed transmutation states otherwise.
            """
        ),
    ),
    (
        "Article Two. Spousal waivers and reimbursement",
        wrap(
            """
            Each spouse waives claims against the other spouse's separate property
            except as provided in the disclosure exhibits. Community contributions
            to separate property create reimbursement rights. The waiver language
            does not waive fiduciary duties created by a later trust, power of
            attorney, or health care directive.
            """
        ),
    ),
    (
        "Article Three. Transfer and gifting limits",
        wrap(
            """
            Gifts, transfers, retitling, and trust funding should not alter
            separate property or community property classification unless the
            transfer document expressly states the intended effect and both spouses
            consent in writing.
            """
        ),
    ),
]

LAW_UPDATE = wrap(
    """
    Synthetic New Law and Attorney Instruction

    Assume the fictional 2026 California Estate Planning Packet Consistency Act
    requires attorney review before an integrated estate packet is finalized.

    Required checks:
    1. Normalize the trust identity across the trust, pour-over will, signing
       memo, and funding instructions.
    2. Flag any mismatch between trustee, executor, health care agent, and
       financial agent order.
    3. Preserve separate property and community property classifications created
       by a prenup before drafting trust funding language.
    4. State whether incapacity is triggered by two physicians, court order,
       inability to make health care decisions, or immediate effectiveness.
    5. Modernize HIPAA and digital health portal language where an AHCD uses
       older privacy wording.

    Client instruction: prepare a packet reconciliation draft that keeps the
    trust as the base document, revises the pour-over trust name/date, adds an
    attorney review flag for the agent-order mismatch, and prevents the funding
    language from accidentally transmuting separate property.
    """
)

SIGNING_MEMO = wrap(
    """
    # Synthetic Signing Memo - Execution Checklist

    This memo intentionally conflicts with the packet so the review tab has
    something to catch.

    - It says the pour-over will refers to the Chen Family Revocable Trust dated
      May 12, 2016, but the uploaded will PDF uses Maya Chen Living Trust dated
      May 21, 2016.
    - It says Daniel Chen is first health care agent, but the AHCD names Priya
      Shah as first health care agent.
    - It says all funded assets are community property, but the prenup preserves
      Solara AI Inc. founder shares as Maya Chen's separate property.
    - It says the POA is springing on physician certification, but the financial
      POA is effective immediately.

    Expected Drafting Magic behavior: mark the matrix stale after this memo is
    pasted into any source card, then regenerate comparison rows that surface
    trust identity, agent order, property character, and incapacity trigger.
    """
)

LEGACY_DOC = wrap(
    """
    SYNTHETIC LEGACY .DOC FORMAT TEST

    This file is plain text saved with a .doc extension. Drafting Magic currently
    accepts .doc in the file picker but does not extract legacy Word binary
    content. Uploading this file should produce a Needs review warning rather
    than silently treating it as usable extracted text.
    """
)


def write_readme(files: list[dict[str, str]]) -> None:
    readme = OUT / "00_README_TESTING_GUIDE.md"
    lines = [
        "# Drafting Magic Synthetic Test Packet",
        "",
        "All documents in this folder are synthetic. They intentionally use fake but realistic names, addresses, and account-like facts so the local privacy filter and Drafting Magic workflow can be tested without real client data.",
        "",
        "## Suggested Upload Mapping",
        "",
        "| Drafting Magic slot | Upload file | Expected behavior |",
        "| --- | --- | --- |",
        "| Revocable living trust | `01_REVOCABLE_LIVING_TRUST__Chen_Family_2016.docx` | Ready; trust identity, trustees, funding, property character, incapacity units |",
        "| Pour-over will | `02_POUR_OVER_WILL__Maya_Chen_Trust_Date_Mismatch.pdf` | Ready; trust name/date mismatch, executor, residue, property silence |",
        "| Advance directive | `03_ADVANCE_HEALTH_CARE_DIRECTIVE__Priya_First_Agent.docx` | Ready; AHCD agent-order mismatch and HIPAA modernization issue |",
        "| Financial POA | `04_DURABLE_FINANCIAL_POA__Daniel_First_Agent.docx` | Ready; effective immediately, financial agent mismatch, gifting limits |",
        "| Prenup | `05_PRENUPTIAL_AGREEMENT__Separate_Property_Constraints.pdf` | Ready; separate/community property, waivers, transfer limits |",
        "",
        "Paste `06_NEW_LAW_AND_ATTORNEY_INSTRUCTION__Packet_Reconciliation.txt` into the **Attorney update or new law** box, then click **Generate comparison**.",
        "",
        "## Extra Negative/Edge Tests",
        "",
        "- `07_SIGNING_MEMO__Conflicting_Execution_Checklist.md`: paste into any source card to test stale matrix regeneration and review flags.",
        "- `08_LEGACY_DOC_FORMAT__Unsupported_Extraction.doc`: upload to confirm unsupported `.doc` extraction becomes **Needs review**.",
        "- `09_SCANNED_PDF_WARNING__Image_Only_Addendum.pdf`: upload to confirm image-only PDFs become **Needs review** until OCR exists.",
        "",
        "## Expected Issues",
        "",
        "- Pour-over will names `Maya Chen Living Trust dated May 21, 2016`; trust document says `Chen Family Revocable Trust dated May 12, 2016`.",
        "- Trust successor trustee order differs from AHCD and financial POA agent order.",
        "- Prenup preserves Solara AI shares as separate property, while the trust funding schedule lacks a property-character legend.",
        "- Trust incapacity requires two physicians or court; POA is effective immediately; AHCD turns on inability to make health care decisions.",
        "- AHCD uses old HIPAA/privacy wording and omits digital health portals.",
        "",
        "## Files",
        "",
    ]
    for item in files:
        lines.append(f"- `{item['file']}` - {item['purpose']}")
    readme.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for path in OUT.glob("*"):
        if path.is_file():
            path.unlink()

    files = [
        {
            "file": "01_REVOCABLE_LIVING_TRUST__Chen_Family_2016.docx",
            "purpose": "Base trust with trustee order, funding schedule, property-character ambiguity, and incapacity trigger.",
        },
        {
            "file": "02_POUR_OVER_WILL__Maya_Chen_Trust_Date_Mismatch.pdf",
            "purpose": "Selectable PDF with pour-over trust identity mismatch.",
        },
        {
            "file": "03_ADVANCE_HEALTH_CARE_DIRECTIVE__Priya_First_Agent.docx",
            "purpose": "AHCD with Priya Shah first, HIPAA modernization issue, and health care decision trigger.",
        },
        {
            "file": "04_DURABLE_FINANCIAL_POA__Daniel_First_Agent.docx",
            "purpose": "Financial POA with Daniel first, immediate effectiveness, gifting and transfer limits.",
        },
        {
            "file": "05_PRENUPTIAL_AGREEMENT__Separate_Property_Constraints.pdf",
            "purpose": "Prenup with separate/community property, spousal waivers, reimbursement, and trust-funding limits.",
        },
        {
            "file": "06_NEW_LAW_AND_ATTORNEY_INSTRUCTION__Packet_Reconciliation.txt",
            "purpose": "Paste into the new law/instruction field.",
        },
        {
            "file": "07_SIGNING_MEMO__Conflicting_Execution_Checklist.md",
            "purpose": "Paste/upload as an edge source to force stale analysis and review flags.",
        },
        {
            "file": "08_LEGACY_DOC_FORMAT__Unsupported_Extraction.doc",
            "purpose": "Unsupported legacy extension warning test.",
        },
        {
            "file": "09_SCANNED_PDF_WARNING__Image_Only_Addendum.pdf",
            "purpose": "Image-only scanned PDF warning test.",
        },
    ]

    make_docx(
        OUT / files[0]["file"],
        "Chen Family Revocable Trust",
        "Synthetic fixture dated May 12, 2016",
        TRUST_BLOCKS,
    )
    make_pdf(
        OUT / files[1]["file"],
        "Pour-Over Will of Maya Chen",
        "Synthetic fixture with intentional trust identity mismatch",
        WILL_BLOCKS,
    )
    make_docx(
        OUT / files[2]["file"],
        "Advance Health Care Directive of Maya Chen",
        "Synthetic fixture naming Priya Shah first health care agent",
        AHCD_BLOCKS,
    )
    make_docx(
        OUT / files[3]["file"],
        "Durable Financial Power of Attorney of Maya Chen",
        "Synthetic fixture naming Daniel Chen first financial agent",
        POA_BLOCKS,
    )
    make_pdf(
        OUT / files[4]["file"],
        "Prenuptial Agreement of Maya Chen and Daniel Chen",
        "Synthetic fixture preserving separate property and transfer limits",
        PRENUP_BLOCKS,
    )
    (OUT / files[5]["file"]).write_text(LAW_UPDATE + "\n", encoding="utf-8")
    (OUT / files[6]["file"]).write_text(SIGNING_MEMO + "\n", encoding="utf-8")
    (OUT / files[7]["file"]).write_text(LEGACY_DOC + "\n", encoding="utf-8")
    make_image_only_pdf(OUT / files[8]["file"])
    write_readme(files)

    manifest = {
        "generated_by": "scripts/generate_drafting_magic_synthetic_docs.py",
        "synthetic": True,
        "fixture_folder": "fixtures/drafting-magic-synthetic-packet",
        "files": files,
        "expected_feature_coverage": [
            "browser DOCX extraction",
            "browser selectable-PDF extraction",
            "TXT/MD paste or upload",
            "unsupported .doc warning",
            "image-only PDF warning",
            "privacy filter tokenization of fake names and addresses",
            "comparison matrix stale/regenerate flow",
            "trust identity mismatch",
            "fiduciary and agent order mismatch",
            "separate property/community property constraints",
            "incapacity trigger mismatch",
            "AHCD/HIPAA modernization issue",
            "draft/review/export path",
        ],
    }
    (OUT / "manifest.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")

    zip_path = OUT.parent / "drafting-magic-synthetic-packet.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for file_path in sorted(OUT.iterdir()):
            if file_path.is_file():
                zf.write(file_path, arcname=f"drafting-magic-synthetic-packet/{file_path.name}")

    print(f"Wrote {OUT}")
    print(f"Wrote {zip_path}")


if __name__ == "__main__":
    main()
